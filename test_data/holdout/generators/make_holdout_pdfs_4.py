# -*- coding: utf-8 -*-
"""4차 홀드아웃: 서로 다른 '문서 형식' 3종.
H = 스캔본처럼 만든 PDF(텍스트 레이어 없음, OCR 경로 강제)
I = DOCX(Word 워드 인제스천 경로)
J = 직접 업로드 JPG(이미지 직접 업로드 경로)
전부 오늘 처음 쓰는 내용(정수기/헬스기구/드론)."""
import fitz  # PyMuPDF
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas as rl_canvas

pdfmetrics.registerFont(UnicodeCIDFont("HYSMyeongJo-Medium"))
styles = getSampleStyleSheet()
title_style = ParagraphStyle("KTitle", parent=styles["Title"], fontName="HYSMyeongJo-Medium", fontSize=18)
h_style = ParagraphStyle("KHeading", parent=styles["Heading2"], fontName="HYSMyeongJo-Medium", fontSize=13)
body_style = ParagraphStyle("KBody", parent=styles["Normal"], fontName="HYSMyeongJo-Medium", fontSize=11, leading=16)


def build_text_pdf(path, blocks):
    doc = SimpleDocTemplate(path, pagesize=A4)
    story = []
    for kind, text in blocks:
        style = {"title": title_style, "h": h_style}.get(kind, body_style)
        story.append(Paragraph(text, style))
        story.append(Spacer(1, 10))
    doc.build(story)


# ---------- H: 스캔본처럼 만들기 (테라클린 정수기) ----------
teraclean_blocks = [
    ("title", "테라클린 정수기 회사소개서"),
    ("h", "TERACLEAN"),
    (
        "b",
        "테라클린은 가정용 정수기를 전문으로 생산하는 제조사입니다. 주력 제품은 TC 시리즈이며, "
        "역삼투압(RO) 필터링 방식을 채택하고 있습니다.",
    ),
    ("h", "RO 필터의 기능"),
    ("b", "1. UV 살균 기능<br/>자외선 램프로 필터를 통과한 물의 잔존 세균을 제거합니다."),
    ("b", "2. 역삼투압 필터링 기능<br/>반투과성 막을 이용해 물속의 중금속과 미세 불순물을 걸러냅니다."),
    ("b", "3. 잔류염소 제거 기능<br/>수돗물에 포함된 잔류염소 냄새를 활성탄 전처리 단계에서 흡착 제거합니다."),
    (
        "b",
        "참고로 본 제품은 역삼투압(RO) 방식을 사용하며, 이는 단순 활성탄 필터링 방식보다 "
        "중금속 제거율이 훨씬 높습니다.",
    ),
    ("h", "정수기 설치 공정"),
    (
        "b",
        "설치 공정은 다음 순서로 진행됩니다.<br/>"
        "1. 벽면 고정<br/>2. 급수라인 연결<br/>3. 필터 장착<br/>4. 누수 테스트<br/>5. 최종 점검",
    ),
]
tmp_pdf = "holdout_H_tmp_native.pdf"
build_text_pdf(tmp_pdf, teraclean_blocks)

# 네이티브 PDF를 고해상도 이미지로 렌더링한 뒤, 이미지만으로 새 PDF를 만든다 (텍스트 레이어 없음 = 스캔본과 동일 취급)
src = fitz.open(tmp_pdf)
c = rl_canvas.Canvas("holdout_H_scanned.pdf", pagesize=A4)
for page in src:
    pix = page.get_pixmap(dpi=200)
    img_path = "holdout_H_page.png"
    pix.save(img_path)
    c.drawImage(ImageReader(img_path), 0, 0, width=A4[0], height=A4[1])
    c.showPage()
c.save()
src.close()
print("wrote holdout_H_scanned.pdf (텍스트 레이어 없음)")

# ---------- I: DOCX (오르카짐 헬스기구) ----------
from docx import Document as DocxDocument
from docx.shared import Pt

d = DocxDocument()
d.add_heading("오르카짐 러닝머신 제품 사양", level=1)
d.add_paragraph(
    "오르카짐은 가정용·상업용 러닝머신을 전문으로 생산하는 회사입니다. "
    "주력 제품은 OG-Run 시리즈이며, 충격흡수 데크 구조를 채택하고 있습니다."
)
d.add_heading("충격흡수 데크의 기능", level=2)
d.add_paragraph("1. 관절 부담 감소 기능: 러닝 시 무릎과 발목에 가해지는 충격을 흡수해 부상 위험을 줄입니다.")
d.add_paragraph("2. 소음 저감 기능: 데크와 프레임 사이 완충재가 운동 중 발생하는 진동 소음을 낮춥니다.")
d.add_paragraph(
    "참고로 본 제품은 유압식 완충 데크를 사용하며, 이는 스프링식 완충 방식보다 내구성이 더 오래 유지됩니다."
)
d.add_heading("모델별 사양", level=2)
table = d.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = "모델명", "최고속도(km/h)", "최대하중(kg)"
rows = [("OG-Run100", "20", "150"), ("OG-Run300", "22", "180"), ("OG-Run500", "25", "200")]
for model, speed, weight in rows:
    cells = table.add_row().cells
    cells[0].text, cells[1].text, cells[2].text = model, speed, weight
d.add_heading("안전 인증", level=2)
d.add_paragraph("OG-Run300 모델의 안전인증번호: SF-GYM-2024-5567")
d.save("holdout_I_orcagym.docx")
print("wrote holdout_I_orcagym.docx")

# ---------- J: 직접 이미지 업로드 (블루윙드론) ----------
bluewing_blocks = [
    ("title", "블루윙드론 회사소개"),
    ("h", "BLUEWING DRONE"),
    (
        "b",
        "블루윙드론은 산업 점검용 드론을 전문으로 제작하는 회사입니다. 주력 제품은 BW-X 시리즈이며, "
        "듀얼 GPS 기반 자동복귀(RTH) 시스템을 탑재하고 있습니다.",
    ),
    ("h", "자동복귀 시스템의 기능"),
    ("b", "1. 신호 유실 시 자동복귀 기능: 조종 신호가 일정 시간 끊기면 이륙 지점으로 자동 복귀합니다."),
    ("b", "2. 저배터리 자동복귀 기능: 배터리 잔량이 20% 이하가 되면 자동으로 복귀를 시작합니다."),
    ("h", "BW-X300 사양"),
    ("b", "최대비행시간: 32분<br/>최대속도: 65km/h<br/>최대통신거리: 8km"),
]
tmp2 = "holdout_J_tmp.pdf"
build_text_pdf(tmp2, bluewing_blocks)
src2 = fitz.open(tmp2)
pix2 = src2[0].get_pixmap(dpi=200)
pix2.save("holdout_J_bluewing.jpg")
src2.close()
print("wrote holdout_J_bluewing.jpg")
