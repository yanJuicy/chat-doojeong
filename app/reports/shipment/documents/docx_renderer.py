"""Default DOCX renderer for a calculated shipment report."""

from io import BytesIO

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .models import ShipmentDocumentView


CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
PASS = "1F3A5F"
WARNING = "7A5A00"


def _set_run_font(
    run,
    *,
    size: float,
    bold: bool = False,
    color: str = "000000",
    name: str = "Calibri",
) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), "Malgun Gothic")


def _set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def _set_paragraph_fill(paragraph, color: str) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    shading = paragraph_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        paragraph_properties.append(shading)
    shading.set(qn("w:fill"), color)


def _set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        margin = tc_mar.find(qn(f"w:{tag}"))
        if margin is None:
            margin = OxmlElement(f"w:{tag}")
            tc_mar.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError("DOCX table widths must total 9360 DXA")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag in ("tblW", "tblInd", "tblLayout"):
        existing = tbl_pr.find(qn(f"w:{tag}"))
        if existing is not None:
            tbl_pr.remove(existing)

    width = OxmlElement("w:tblW")
    width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    width.set(qn("w:type"), "dxa")
    tbl_pr.append(width)
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for column_width in widths:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(column_width))
        grid.append(grid_column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            cell_width = tc_pr.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                tc_pr.append(cell_width)
            cell_width.set(qn("w:w"), str(widths[index]))
            cell_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _repeat_header(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    row_properties.append(marker)


def _set_cell_text(
    cell,
    text: str,
    *,
    size: float = 9,
    bold: bool = False,
    color: str = "000000",
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    _set_run_font(paragraph.add_run(text), size=size, bold=bold, color=color)


def _add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def _add_table_spacing(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(paragraph.add_run("Page "), size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


class ShipmentDocxRenderer:
    """Render the default, customer-neutral daily shipment report."""

    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    extension = "docx"

    def render(self, view: ShipmentDocumentView) -> bytes:
        document = Document()
        self._configure_document(document)
        self._add_masthead(document, view)
        self._add_summary(document, view)
        self._add_metrics(document, view)
        self._add_details(document, view)
        self._add_exceptions(document, view)
        self._add_validations(document, view)

        document.core_properties.title = view.title
        document.core_properties.subject = "일일 출하보고서"
        document.core_properties.author = "Shipment Report API"
        stream = BytesIO()
        document.save(stream)
        return stream.getvalue()

    @staticmethod
    def _configure_document(document: Document) -> None:
        section = document.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        normal = document.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.25
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

        title = document.styles["Title"]
        title.font.name = "Calibri"
        title.font.size = Pt(24)
        title.font.bold = True
        title.font.color.rgb = RGBColor.from_string(INK)
        title.paragraph_format.space_before = Pt(0)
        title.paragraph_format.space_after = Pt(6)
        title._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

        subtitle = document.styles["Subtitle"]
        subtitle.font.name = "Calibri"
        subtitle.font.size = Pt(12)
        subtitle.font.color.rgb = RGBColor.from_string(MUTED)
        subtitle.paragraph_format.space_after = Pt(14)
        subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

        heading_tokens = (
            ("Heading 1", 16, BLUE, 18, 10),
            ("Heading 2", 13, BLUE, 14, 7),
            ("Heading 3", 12, DARK_BLUE, 10, 5),
        )
        for name, size, color, before, after in heading_tokens:
            style = document.styles[name]
            style.font.name = "Calibri"
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_run_font(header.add_run("DAILY SHIPMENT REPORT"), size=9, bold=True, color=MUTED)
        _add_page_number(section.footer.paragraphs[0])

    @staticmethod
    def _add_masthead(document: Document, view: ShipmentDocumentView) -> None:
        kicker = document.add_paragraph()
        kicker.paragraph_format.space_after = Pt(3)
        _set_run_font(kicker.add_run("SHIPMENT OPERATIONS"), size=9.5, bold=True, color=BLUE)
        document.add_paragraph(view.title, style="Title")
        document.add_paragraph(
            f"{view.customer_name} · {view.delivery_location}",
            style="Subtitle",
        )

        rows = (
            ("보고일자", view.report_date),
            ("고객사", f"{view.customer_name} ({view.customer_code})"),
            ("납품장소", view.delivery_location),
            ("보고서 ID", view.report_id),
        )
        for label, value in rows:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.15
            _set_run_font(paragraph.add_run(f"{label}: "), size=9.5, bold=True, color=DARK_BLUE)
            _set_run_font(paragraph.add_run(value), size=9.5)

    @staticmethod
    def _add_summary(document: Document, view: ShipmentDocumentView) -> None:
        _add_heading(document, "1. 출하 요약")
        callout = document.add_paragraph()
        callout.paragraph_format.left_indent = Inches(0.08)
        callout.paragraph_format.right_indent = Inches(0.08)
        callout.paragraph_format.space_before = Pt(4)
        callout.paragraph_format.space_after = Pt(8)
        callout.paragraph_format.line_spacing = 1.25
        _set_paragraph_fill(callout, LIGHT_FILL)
        _set_run_font(callout.add_run(view.summary), size=10.5, color=INK)

    @staticmethod
    def _add_metrics(document: Document, view: ShipmentDocumentView) -> None:
        _add_heading(document, "2. 종합 현황")
        table = document.add_table(rows=2, cols=len(view.metrics))
        table.style = "Table Grid"
        _repeat_header(table.rows[0])
        column_widths = [1872] * len(view.metrics)
        for index, metric in enumerate(view.metrics):
            _set_cell_fill(table.cell(0, index), HEADER_FILL)
            _set_cell_text(
                table.cell(0, index), metric.label, size=9, bold=True, color=DARK_BLUE,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            _set_cell_text(
                table.cell(1, index), metric.value, size=10.5, bold=True, color=INK,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
        _set_table_geometry(table, column_widths)
        _add_table_spacing(document)

    @staticmethod
    def _add_details(document: Document, view: ShipmentDocumentView) -> None:
        _add_heading(document, "3. 품목별 출하 현황")
        headers = ("품목", "계획", "출하", "미출하", "초과", "달성률", "상태", "사유")
        widths = [1900, 1050, 1050, 950, 950, 850, 1000, 1610]
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        _repeat_header(table.rows[0])
        for index, header in enumerate(headers):
            _set_cell_fill(table.cell(0, index), HEADER_FILL)
            _set_cell_text(
                table.cell(0, index), header, size=8.5, bold=True, color=DARK_BLUE,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
        for line in view.details:
            cells = table.add_row().cells
            values = (
                f"{line.item_code}\n{line.item_name}",
                line.planned_qty,
                line.shipped_qty,
                line.unshipped_qty,
                line.over_shipped_qty,
                line.achievement_rate,
                line.status,
                line.note,
            )
            for index, value in enumerate(values):
                color = WARNING if index == 6 and line.status != "정상" else "000000"
                _set_cell_text(
                    cells[index], value, size=8.5, bold=index == 6, color=color,
                    align=WD_ALIGN_PARAGRAPH.CENTER if index not in (0, 7) else WD_ALIGN_PARAGRAPH.LEFT,
                )
        _set_table_geometry(table, widths)
        _add_table_spacing(document)

    @staticmethod
    def _add_exceptions(document: Document, view: ShipmentDocumentView) -> None:
        _add_heading(document, "4. 특이사항 및 경고")
        messages = view.warnings or ("확인 필요한 경고가 없습니다.",)
        for message in messages:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.08)
            paragraph.paragraph_format.right_indent = Inches(0.08)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.15
            _set_paragraph_fill(paragraph, LIGHT_FILL)
            _set_run_font(
                paragraph.add_run(message), size=9.5,
                color=WARNING if view.warnings else PASS,
            )

    @staticmethod
    def _add_validations(document: Document, view: ShipmentDocumentView) -> None:
        _add_heading(document, "5. 데이터 검증")
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        headers = ("검증 코드", "결과", "메시지")
        for index, header in enumerate(headers):
            _set_cell_fill(table.cell(0, index), HEADER_FILL)
            _set_cell_text(
                table.cell(0, index), header, size=9, bold=True, color=DARK_BLUE,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
        _repeat_header(table.rows[0])
        for validation in view.validations:
            cells = table.add_row().cells
            _set_cell_text(cells[0], validation.code, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(
                cells[1], validation.status, size=9, bold=True,
                color=PASS if validation.status == "PASS" else WARNING,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            _set_cell_text(cells[2], validation.message, size=9)
        _set_table_geometry(table, [1900, 1100, 6360])
        _add_table_spacing(document)

        generated = document.add_paragraph()
        generated.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        generated.paragraph_format.space_before = Pt(4)
        generated.paragraph_format.space_after = Pt(0)
        _set_run_font(
            generated.add_run(f"생성시각(UTC): {view.generated_at}"),
            size=8.5,
            color=MUTED,
        )
