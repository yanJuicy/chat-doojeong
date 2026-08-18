"""버전형 DOCX 양식에 주간보고서 데이터를 채우는 렌더러."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from ....db.models import WorkItemStatus
from ..models import WeeklyReportDraft


CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CURRENT_WIDTHS = [1300, 2700, 1200, 4160]
NEXT_WIDTHS = [1300, 3700, 1200, 3160]
HEADER_FILL = "F2F4F7"
DARK_BLUE = "1F4D78"
MUTED = "667085"
WARNING = "7A5A00"

_STATUS_LABELS = {
    WorkItemStatus.PLANNED: "예정",
    WorkItemStatus.IN_PROGRESS: "진행 중",
    WorkItemStatus.COMPLETED: "완료",
    WorkItemStatus.ON_HOLD: "보류",
}


def _set_run_font(run, *, size: float, bold: bool = False, color: str = "000000") -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Malgun Gothic")


def _set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), color)


def _set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        margin = tc_mar.find(qn(f"w:{tag}"))
        if margin is None:
            margin = OxmlElement(f"w:{tag}")
            tc_mar.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError("DOCX table widths must total 9360 DXA")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag in ("tblW", "tblInd", "tblLayout"):
        existing = tbl_pr.find(qn(f"w:{tag}"))
        if existing is not None:
            tbl_pr.remove(existing)
    width = OxmlElement("w:tblW")
    width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    width.set(qn("w:type"), "dxa")
    tbl_pr.append(width)
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            cell_width = tc_pr.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                tc_pr.append(cell_width)
            cell_width.set(qn("w:w"), str(widths[index]))
            cell_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:tblHeader")) is None:
        marker = OxmlElement("w:tblHeader")
        marker.set(qn("w:val"), "true")
        properties.append(marker)


def _set_cell_text(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.10
    _set_run_font(paragraph.add_run(text), size=9, bold=bold)


def _find_table(document: Document, marker: str):
    for table in document.tables:
        if any(marker in cell.text for row in table.rows for cell in row.cells):
            return table
    raise ValueError(f"DOCX template marker not found: {marker}")


def _remove_marker_row(table, marker: str) -> None:
    for row in list(table.rows):
        if any(marker in cell.text for cell in row.cells):
            table._tbl.remove(row._tr)
            return
    raise ValueError(f"DOCX template row marker not found: {marker}")


def _replace_placeholders(document: Document, values: dict[str, str]) -> None:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        original = paragraph.text
        updated = original
        for marker, value in values.items():
            updated = updated.replace(marker, value)
        if updated != original:
            paragraph.text = ""
            _set_run_font(paragraph.add_run(updated), size=9.5)


class WeeklyDocxRenderer:
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def render(self, report: WeeklyReportDraft, template_path: Path) -> bytes:
        if not template_path.is_file() or template_path.suffix.lower() != ".docx":
            raise ValueError(f"유효한 DOCX 양식 파일이 아닙니다: {template_path}")
        document = Document(str(template_path))
        _replace_placeholders(
            document,
            {
                "{{author}}": report.author or "-",
                "{{department}}": report.department or "-",
                "{{period}}": f"{report.period_start.isoformat()} ~ {report.period_end.isoformat()}",
                "{{cutoff_date}}": report.cutoff_date.isoformat(),
                "{{next_period}}": f"{report.next_week_start.isoformat()} ~ {report.next_week_end.isoformat()}",
                "{{warnings}}": "\n".join(report.warnings) if report.warnings else "검토가 필요한 경고가 없습니다.",
            },
        )

        current_table = _find_table(document, "{{current_category}}")
        _remove_marker_row(current_table, "{{current_category}}")
        for item in report.current_week:
            details = list(item.activity_details)
            if item.result and item.result not in details:
                details.append(item.result)
            values = (
                item.category or "-",
                item.title,
                _STATUS_LABELS[item.status],
                "\n".join(details) if details else "-",
            )
            cells = current_table.add_row().cells
            for index, value in enumerate(values):
                _set_cell_text(cells[index], value, bold=index == 2)
        if not report.current_week:
            cells = current_table.add_row().cells
            for index, value in enumerate(("-", "등록된 금주 진행사항이 없습니다.", "-", "-")):
                _set_cell_text(cells[index], value)
        _repeat_header(current_table.rows[0])
        for cell in current_table.rows[0].cells:
            _set_cell_fill(cell, HEADER_FILL)
        _set_table_geometry(current_table, CURRENT_WIDTHS)

        next_table = _find_table(document, "{{next_category}}")
        _remove_marker_row(next_table, "{{next_category}}")
        for item in report.next_week:
            values = (
                item.category or "-",
                item.plan,
                item.target_date.isoformat() if item.target_date else "-",
                "이월" if item.carry_over else "-",
            )
            cells = next_table.add_row().cells
            for index, value in enumerate(values):
                _set_cell_text(cells[index], value, bold=index == 3 and item.carry_over)
        if not report.next_week:
            cells = next_table.add_row().cells
            for index, value in enumerate(("-", "등록된 차주 진행 계획이 없습니다.", "-", "-")):
                _set_cell_text(cells[index], value)
        _repeat_header(next_table.rows[0])
        for cell in next_table.rows[0].cells:
            _set_cell_fill(cell, HEADER_FILL)
        _set_table_geometry(next_table, NEXT_WIDTHS)

        document.core_properties.title = "주간 업무보고서"
        document.core_properties.subject = "금주 진행사항 및 차주 진행 계획"
        document.core_properties.author = "Weekly Report API"
        stream = BytesIO()
        document.save(stream)
        return stream.getvalue()
