"""
WeeklyReportView를 "주간 업무실적 및 계획" 양식의 DOCX로 렌더링한다. 구분 행 개수는
view.rows 길이만큼 동적으로 생성된다(원본 문서가 구분을 여러 개 썼으면 그대로 여러 행).
"""
from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

# 원본 "주간 업무실적 및 계획" 양식의 열 비율(좁은 "구분" 칸 + 넓은 실적/계획 칸 두 개)을
# 그대로 따른다 — python-docx는 열 너비를 지정 안 하면 3등분으로 균등 배분해서, "구분"
# 칸만 유독 헐렁하고 내용 칸은 문단 텍스트가 부대끼며 잡히는 문제가 있었다.
# 구분 칸은 "사업관리"(4글자)만 가정했던 예전과 달리 이제 "인프라 / 배포"처럼 더 긴
# 구분명도 들어와서, 살짝 더 넓게 잡는다(줄바꿈 허용).
_CATEGORY_COLUMN_WIDTH = Inches(1.3)
_CONTENT_COLUMN_WIDTH = Inches(2.8)

from .weekly_report_composer import ReportPeriod, WeeklyReportView, week_of_month_label


def _set_run_font(run, *, size: float, bold: bool = False) -> None:
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    fonts.set(qn("w:eastAsia"), "맑은 고딕")


def _set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    tc_pr.append(shading)


def _set_cell_margins(cell, *, top: int = 120, bottom: int = 120, left: int = 140, right: int = 140) -> None:
    """셀 안쪽 여백(twips, 1/1440인치)을 넉넉히 줘서 글자가 테두리에 바짝 붙지 않게 한다."""
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tc_pr.append(margins)


def _set_cell_width(cell, width) -> None:
    """python-docx는 열 너비를 table.columns뿐 아니라 각 셀에도 따로 지정해야 워드에서 반영된다."""
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(width.twips))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def _format_period(period: ReportPeriod) -> str:
    return f"({period.period_start.strftime('%Y.%m.%d')} ~ {period.period_end.strftime('%Y.%m.%d')})"


def _bullet_prefix(source_format: str | None) -> str:
    """부서가 업로드한 원본 문서가 쓰던 글머리 기호를 그대로 재현한다. 원본이 문장 하나로
    (기호 없이) 쓰여 있었거나(prose), 문서를 한 번도 안 올린 부서(None)는 기본값 "•"."""
    if source_format and source_format.startswith("bullet:"):
        return source_format.removeprefix("bullet:")
    return "•"


def _fill_items_cell(cell, items: list, *, source_format: str | None = None, size: float = 10) -> None:
    paragraph = cell.paragraphs[0]
    if not items:
        _set_run_font(paragraph.add_run("-"), size=size)
        return
    if source_format == "prose":
        # 원본 문서가 글머리 기호 없이 문장 하나로 쓰던 부서 — 항목을 이어붙여 같은 형태로 낸다.
        joined = " ".join(item.content for item in items)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        _set_run_font(paragraph.add_run(joined), size=size)
        return
    bullet = _bullet_prefix(source_format)
    for index, item in enumerate(items):
        target = paragraph if index == 0 else cell.add_paragraph()
        target.paragraph_format.space_before = Pt(0)
        target.paragraph_format.space_after = Pt(2)
        _set_run_font(target.add_run(f"{bullet} {item.content}"), size=size)


def render_weekly_report_docx(view: WeeklyReportView) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    week_label = week_of_month_label(view.current_period.period_start)
    _set_run_font(title.add_run(f"주간 업무실적 및 계획 ({week_label})"), size=18, bold=True)

    dept_paragraph = document.add_paragraph()
    dept_paragraph.paragraph_format.space_after = Pt(10)
    _set_run_font(dept_paragraph.add_run(f"■ 부서명 : {view.department}"), size=10.5, bold=True)

    table = document.add_table(rows=1 + len(view.rows), cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False  # 열 너비를 아래에서 직접 지정하므로, 워드가 내용 길이 보고 임의로 재배분하지 않게 끈다.
    column_widths = (_CATEGORY_COLUMN_WIDTH, _CONTENT_COLUMN_WIDTH, _CONTENT_COLUMN_WIDTH)
    for row in table.rows:
        for cell, width in zip(row.cells, column_widths):
            _set_cell_width(cell, width)
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    headers = (
        "구분",
        f"업무 실적\n{_format_period(view.current_period)}",
        f"업무 계획\n{_format_period(view.next_period)}",
    )
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        _set_cell_fill(cell, "D9E2F3")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lines = header.split("\n")
        _set_run_font(paragraph.add_run(lines[0]), size=10.5, bold=True)
        for line in lines[1:]:
            sub = cell.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_run_font(sub.add_run(line), size=9)

    for row_index, category_row in enumerate(view.rows, start=1):
        row = table.rows[row_index]
        category_paragraph = row.cells[0].paragraphs[0]
        category_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(category_paragraph.add_run(category_row.category), size=10.5, bold=True)

        _fill_items_cell(row.cells[1], category_row.current_items, source_format=view.source_format)
        _fill_items_cell(row.cells[2], category_row.next_items, source_format=view.source_format)

    document.core_properties.title = f"주간 업무실적 및 계획 ({view.department})"
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()
