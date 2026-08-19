# -*- coding: utf-8 -*-
"""
scripts/seed_weekly_reports.py (가칭)
=======================================
[B. 합성 데이터 + 적재] 담당 산출물.

역할:
  1) 1~6월(26주치) "시군 특화 일자리 사업단" 가상 주간보고서를 실제 회사 양식(표 구조)
     그대로 docx 26개로 생성한다.
  2) (선택) 생성된 docx를 기존 문서 업로드 API(POST /api/v1/upload)로 라벨 "주간보고서"를
     붙여 일괄 적재한다. — 이 부분은 실제 서버가 떠 있는 로컬 환경(mac)에서만 동작한다.

사용법 (로컬 mac, .venv 활성화 후):
    python scripts/seed_weekly_reports.py                # docx만 생성 (기본)
    python scripts/seed_weekly_reports.py --upload        # 생성 + 서버 업로드까지
    python scripts/seed_weekly_reports.py --upload --host http://127.0.0.1:8000
"""
import argparse
import datetime
import os
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# 0. 기본 설정 — 필요하면 이 값만 바꿔서 재생성하면 된다.
# ---------------------------------------------------------------------------
DEPARTMENT_NAME = "일자리정책팀"   # TODO: 실제 부서명으로 교체
PROJECT_NAME = "시군 특화 일자리 사업단"
CATEGORY = "사업관리"
OUTPUT_DIR = "synthetic_weekly_reports"
LABEL = "주간보고서"

# ---------------------------------------------------------------------------
# 1. 26주치 실적/계획 데이터 — 1월 첫째주 ~ 6월 말, 6개월 사업 진행 흐름 반영
#    (1월: 계획수립 → 2월: 설명회·모집 → 3월: 매칭 개시 → 4월: 채용박람회
#     → 5월: 중간점검·홍보 → 6월: 상반기 결산·하반기 계획)
# ---------------------------------------------------------------------------
WEEKLY_CONTENT = [
    {
        "실적": ["2026년 시군 특화 일자리 사업단 운영계획(안) 수립", "사업 예산(안) 검토 및 소관 부서 협의", "참여 대상 시군 후보 목록 정리"],
        "계획": ["사업 예산 확정", "참여 시군 대상 사업설명 자료 초안 작성"],
    },
    {
        "실적": ["사업 예산 확정", "참여 시군 대상 사업설명 자료 초안 작성 완료"],
        "계획": ["시군 실무협의체 구성", "사업설명회 일정 조율"],
    },
    {
        "실적": ["시군 실무협의체 1차 구성 완료(3개 시군)", "사업설명회 일정 조율 완료"],
        "계획": ["사업설명회 개최 준비(장소 섭외, 참석대상 안내)"],
    },
    {
        "실적": ["사업설명회 개최 준비 완료", "참여기업 모집공고(안) 마련"],
        "계획": ["참여 시군 대상 사업설명회 개최", "참여기업 모집공고 게시"],
    },
    {
        "실적": ["참여 시군 대상 사업설명회 개최(3개 시군, 참석 42명)", "참여기업 모집공고 게시"],
        "계획": ["참여기업 접수 개시", "취업상담인력 채용공고 게시"],
    },
    {
        "실적": ["참여기업 접수 개시(1주차 12개사 접수)", "취업상담인력 채용공고 게시"],
        "계획": ["상담인력 서류·면접 전형 진행"],
    },
    {
        "실적": ["상담인력 서류·면접 전형 완료", "참여기업 접수 마감(누적 27개사)"],
        "계획": ["상담인력 채용 확정 및 시군별 배치", "참여기업 명단 확정"],
    },
    {
        "실적": ["상담인력 3명 채용 확정 및 시군별 배치", "참여기업 명단 확정(27개사)"],
        "계획": ["구직자 모집 홍보물 제작", "구인·구직 매칭시스템 시범 운영 준비"],
    },
    {
        "실적": ["구직자 모집 홍보물 제작 완료", "구인·구직 매칭시스템 시범 운영 점검"],
        "계획": ["구직자 모집 개시", "시군별 취업상담창구 운영 개시"],
    },
    {
        "실적": ["구직자 모집 개시(1주차 접수 58명)", "시군별 취업상담창구 3개소 운영 개시"],
        "계획": ["1차 구인·구직 매칭 진행"],
    },
    {
        "실적": ["1차 구인·구직 매칭 진행(38건 매칭)", "상담창구 운영 현황 점검"],
        "계획": ["1차 채용연계 결과 취합 및 참여기업 애로사항 조사"],
    },
    {
        "실적": ["1차 채용연계 결과 취합(최종 채용 9명)", "참여기업 애로사항 조사 실시"],
        "계획": ["애로사항 개선방안 마련", "2차 구직자 모집 준비"],
    },
    {
        "실적": ["애로사항 개선방안 마련(매칭 조건 세분화 등)", "2차 구직자 모집 준비 완료"],
        "계획": ["채용박람회 개최계획 수립"],
    },
    {
        "실적": ["채용박람회 개최계획 수립", "참여기업 대상 박람회 참가 수요조사"],
        "계획": ["채용박람회 장소·일정 확정", "참가기업 확정"],
    },
    {
        "실적": ["채용박람회 장소·일정 확정(4월 넷째주)", "참가기업 18개사 확정"],
        "계획": ["채용박람회 홍보 및 구직자 사전등록 진행"],
    },
    {
        "실적": ["채용박람회 홍보 실시", "구직자 사전등록 진행(97명)"],
        "계획": ["시군 특화 채용박람회 개최"],
    },
    {
        "실적": ["시군 특화 채용박람회 개최(참가기업 18개사, 구직자 112명 참석)", "현장면접 진행"],
        "계획": ["박람회 현장면접 결과 취합 및 후속 채용연계"],
    },
    {
        "실적": ["박람회 현장면접 결과 취합(현장채용 14명)", "후속 채용연계 진행"],
        "계획": ["상반기 중간 성과 점검"],
    },
    {
        "실적": ["상반기 중간 성과 점검 실시(누적 채용연계 23명)", "시군별 운영현황 비교분석"],
        "계획": ["우수사례 발굴 및 정리"],
    },
    {
        "실적": ["우수사례 발굴(2개 시군 우수사례 선정)", "사업 홍보자료 제작 착수"],
        "계획": ["지역 언론 대상 사업 홍보 추진"],
    },
    {
        "실적": ["지역 언론 대상 사업 홍보 추진(보도자료 배포)", "2차 참여기업 추가모집 개시"],
        "계획": ["2차 참여기업 접수 및 심사"],
    },
    {
        "실적": ["2차 참여기업 접수 마감(9개사 추가)", "접수기업 심사 실시"],
        "계획": ["상반기 사업 실적 결산 착수"],
    },
    {
        "실적": ["상반기 사업 실적 결산 착수", "시군별 실적 데이터 취합"],
        "계획": ["상반기 결산 보고서(안) 작성"],
    },
    {
        "실적": ["상반기 결산 보고서(안) 작성 완료", "내부 검토 실시"],
        "계획": ["하반기 사업 추진계획(안) 수립"],
    },
    {
        "실적": ["하반기 사업 추진계획(안) 수립", "참여 시군 확대(안) 검토"],
        "계획": ["상반기 성과보고회 준비"],
    },
    {
        "실적": ["상반기 성과보고회 개최(참여기관 관계자 대상, 누적 채용연계 37명 실적 보고)", "하반기 계획 공유"],
        "계획": ["하반기 사업 개시", "3차 참여기업 모집"],
    },
]

assert len(WEEKLY_CONTENT) == 26, f"26주 데이터가 필요한데 {len(WEEKLY_CONTENT)}개 있음"


def get_week_ranges(n_weeks=26, start_year=2026):
    """1월 첫째 월요일부터 시작하는 월~금 기준 n주 (start, end) 리스트."""
    d = datetime.date(start_year, 1, 1)
    while d.weekday() != 0:  # 0 = Monday
        d += datetime.timedelta(days=1)
    ranges = []
    for _ in range(n_weeks):
        ranges.append((d, d + datetime.timedelta(days=4)))
        d += datetime.timedelta(days=7)
    return ranges


def fmt(d: datetime.date) -> str:
    return f"{d.year}. {d.month}. {d.day}."


def set_cell_border(cell, **kwargs):
    """셀 테두리를 얇은 실선으로 지정 (python-docx는 기본 API가 없어 XML 직접 조작)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        el = OxmlElement(tag)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def build_docx(week_no: int, start: datetime.date, end: datetime.date, next_start, next_end, content: dict, out_path: str):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("주간 업무실적 및 계획")
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph()

    # 부서명
    dept = doc.add_paragraph()
    r = dept.add_run(f"■ 부서명: {DEPARTMENT_NAME}")
    r.bold = True
    r.font.size = Pt(12)

    doc.add_paragraph()

    # 표: 헤더 1행 + 데이터 1행, 3열
    table = doc.add_table(rows=2, cols=3)
    table.autofit = False
    widths = [Cm(2.2), Cm(8.0), Cm(8.0)]
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w

    header_cells = table.rows[0].cells
    headers = ["구분", f"업무 실적 ({fmt(start)} ~ {fmt(end)})", f"업무 계획 ({fmt(next_start)} ~ {fmt(next_end)})"]
    for cell, text in zip(header_cells, headers):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(11)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell)

    data_cells = table.rows[1].cells

    # 구분 셀
    c0 = data_cells[0]
    c0.text = ""
    p = c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(CATEGORY)
    r.font.size = Pt(11)
    c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(c0)

    def fill_content_cell(cell, items):
        cell.text = ""
        p0 = cell.paragraphs[0]
        r0 = p0.add_run(PROJECT_NAME)
        r0.bold = True
        r0.font.size = Pt(10.5)
        for item in items:
            p = cell.add_paragraph()
            p.paragraph_format.left_indent = Pt(10)
            r = p.add_run(f"- {item}")
            r.font.size = Pt(10.5)
        set_cell_border(cell)

    fill_content_cell(data_cells[1], content["실적"])
    fill_content_cell(data_cells[2], content["계획"])

    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--upload", action="store_true", help="생성 후 서버에 업로드까지 진행")
    parser.add_argument("--host", default="http://127.0.0.1:8000", help="백엔드 서버 주소")
    parser.add_argument("--outdir", default=OUTPUT_DIR)
    args = parser.parse_args()

    os.makedirs(args.outdir, exist_ok=True)
    ranges = get_week_ranges(26)

    files = []
    for i, ((start, end), content) in enumerate(zip(ranges, WEEKLY_CONTENT), start=1):
        next_start, next_end = ranges[i] if i < len(ranges) else (end + datetime.timedelta(days=3), end + datetime.timedelta(days=7))
        fname = f"주간업무보고_{start.month:02d}월_{i:02d}주차_{start.strftime('%Y%m%d')}.docx"
        out_path = os.path.join(args.outdir, fname)
        build_docx(i, start, end, next_start, next_end, content, out_path)
        files.append(out_path)
        print(f"[{i:02d}/26] 생성 완료: {fname}")

    print(f"\n총 {len(files)}개 문서 생성 완료 → {args.outdir}/")

    if args.upload:
        upload_all(files, args.host)


def upload_all(files, host):
    """기존 업로드 API로 라벨 '주간보고서'를 붙여 일괄 적재.
    실제 로컬 서버(uvicorn)가 떠 있어야 동작한다 — 이 저장소 안에서 직접 실행할 것.
    """
    try:
        import requests
    except ImportError:
        print("requests 패키지가 없습니다: pip install requests --break-system-packages")
        sys.exit(1)

    ok, fail = 0, 0
    for path in files:
        fname = os.path.basename(path)
        with open(path, "rb") as f:
            files_payload = {"file": (fname, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            data = {"labels": LABEL}
            try:
                resp = requests.post(f"{host}/api/v1/upload", files=files_payload, data=data, timeout=30)
                if resp.status_code == 200:
                    print(f"업로드 성공: {fname}")
                    ok += 1
                else:
                    print(f"업로드 실패({resp.status_code}): {fname} — {resp.text[:200]}")
                    fail += 1
            except Exception as e:
                print(f"업로드 예외: {fname} — {e}")
                fail += 1

    print(f"\n업로드 결과: 성공 {ok} / 실패 {fail}")
    if fail:
        print("실패한 파일은 UI(문서 관리 > 새 문서 등록)에서 라벨 '주간보고서'로 수동 업로드하세요.")


if __name__ == "__main__":
    main()
