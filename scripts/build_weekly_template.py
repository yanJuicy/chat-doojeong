"""주간 업무보고서 기본 DOCX 양식을 재현 가능하게 생성한다."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
HEADER_FILL = "F2F4F7"


def set_run_font(run, *, size: float, bold: bool = False, color: str = "000000") -> None:
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Malgun Gothic")


def set_cell_fill(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for tag, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        margin = margins.find(qn(f"w:{tag}"))
        if margin is None:
            margin = OxmlElement(f"w:{tag}")
            margins.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError("table widths must total 9360 DXA")
    table.autofit = False
    properties = table._tbl.tblPr
    for tag in ("tblW", "tblInd", "tblLayout"):
        existing = properties.find(qn(f"w:{tag}"))
        if existing is not None:
            properties.remove(existing)
    width = OxmlElement("w:tblW")
    width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    width.set(qn("w:type"), "dxa")
    properties.append(width)
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    indent.set(qn("w:type"), "dxa")
    properties.append(indent)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    properties.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            properties = cell._tc.get_or_add_tcPr()
            cell_width = properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                properties.append(cell_width)
            cell_width.set(qn("w:w"), str(widths[index]))
            cell_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = "000000") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.10
    set_run_font(paragraph.add_run(text), size=9, bold=bold, color=color)


def repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(paragraph.add_run("Page "), size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    title = document.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(23)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    subtitle = document.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.add_run("WEEKLY WORK REPORT"), size=9, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_header_table(document: Document) -> None:
    table = document.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    values = (
        ("작성자", "{{author}}", "부서", "{{department}}"),
        ("보고 기간", "{{period}}", "작성 기준일", "{{cutoff_date}}"),
    )
    for row_index, row_values in enumerate(values):
        for column_index, value in enumerate(row_values):
            label = column_index % 2 == 0
            set_cell_text(table.cell(row_index, column_index), value, bold=label, color=DARK_BLUE if label else "000000")
            if label:
                set_cell_fill(table.cell(row_index, column_index), HEADER_FILL)
    set_table_geometry(table, [1300, 3380, 1300, 3380])


def add_data_table(document: Document, headers: tuple[str, ...], markers: tuple[str, ...], widths: list[int]) -> None:
    table = document.add_table(rows=2, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.cell(0, index), header, bold=True, color=DARK_BLUE)
        set_cell_fill(table.cell(0, index), HEADER_FILL)
        set_cell_text(table.cell(1, index), markers[index])
    repeat_header(table.rows[0])
    set_table_geometry(table, widths)


def build(output: Path) -> None:
    document = Document()
    configure_document(document)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    set_run_font(kicker.add_run("WEEKLY STATUS"), size=9.5, bold=True, color=BLUE)
    document.add_paragraph("주간 업무보고서", style="Title")
    document.add_paragraph("금주 진행사항 및 차주 진행 계획", style="Subtitle")
    add_header_table(document)

    document.add_heading("1. 금주 진행사항", level=1)
    add_data_table(
        document,
        ("업무 구분", "업무", "상태", "진행 결과 및 특이사항"),
        ("{{current_category}}", "{{current_title}}", "{{current_status}}", "{{current_result}}"),
        [1300, 2700, 1200, 4160],
    )

    document.add_heading("2. 차주 진행 계획", level=1)
    period = document.add_paragraph()
    period.paragraph_format.space_after = Pt(6)
    set_run_font(period.add_run("계획 기간: "), size=9.5, bold=True, color=DARK_BLUE)
    set_run_font(period.add_run("{{next_period}}"), size=9.5)
    add_data_table(
        document,
        ("업무 구분", "계획 내용", "목표일", "비고"),
        ("{{next_category}}", "{{next_plan}}", "{{next_target}}", "{{next_note}}"),
        [1300, 3700, 1200, 3160],
    )

    document.add_heading("3. 검토 사항", level=1)
    warnings = document.add_paragraph("{{warnings}}")
    warnings.paragraph_format.space_before = Pt(0)
    warnings.paragraph_format.space_after = Pt(0)
    warnings.paragraph_format.line_spacing = 1.10

    document.core_properties.title = "주간 업무보고서 기본 양식"
    document.core_properties.subject = "금주 진행사항 및 차주 진행 계획"
    document.core_properties.author = "Weekly Report API"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("app/report_templates/weekly/default_v1.docx"),
    )
    args = parser.parse_args()
    build(args.output)


if __name__ == "__main__":
    main()
