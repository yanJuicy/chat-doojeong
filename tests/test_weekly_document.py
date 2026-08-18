import tempfile
import unittest
from datetime import date, datetime, timezone
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.db.models import GeneratedReport, ReportTemplate, WorkItemStatus
from app.reports.weekly import WeeklyReportDraft
from app.reports.weekly.documents import WeeklyDocumentService, WeeklyDocxRenderer
from app.reports.weekly.models import WeeklyPlanItem, WeeklyProgressItem


PROJECT_ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = PROJECT_ROOT / "app" / "report_templates" / "weekly" / "default_v1.docx"


def sample_report() -> WeeklyReportDraft:
    return WeeklyReportDraft(
        period_start=date(2026, 8, 17),
        period_end=date(2026, 8, 21),
        cutoff_date=date(2026, 8, 19),
        next_week_start=date(2026, 8, 24),
        next_week_end=date(2026, 8, 28),
        author="홍길동",
        department="개발팀",
        current_week=[
            WeeklyProgressItem(
                work_item_id="current-1",
                title="주간보고서 DOCX 생성 기능 개발",
                category="개발",
                status=WorkItemStatus.COMPLETED,
                activity_details=["버전형 양식을 적용하고 다운로드 API를 구현함"],
                result="자동화 테스트 통과",
                completed_on=date(2026, 8, 19),
            )
        ],
        next_week=[
            WeeklyPlanItem(
                work_item_id="next-1",
                title="운영 검증",
                category="검증",
                plan="실제 사용자 데이터로 주간보고서 생성 결과 확인",
                target_date=date(2026, 8, 25),
                carry_over=False,
                reasons=["NEXT_WEEK_DUE"],
            )
        ],
        warnings=["금주 진행사항 1건을 검토하세요."],
        generated_at=datetime(2026, 8, 19, 6, 0, tzinfo=timezone.utc),
    )


class FakeDocumentRepository:
    def __init__(self) -> None:
        self.template = ReportTemplate(
            id="weekly-default-v1",
            report_type="WEEKLY",
            name="기본 주간보고서",
            version=1,
            file_path=str(TEMPLATE_PATH),
            is_active=True,
        )
        self.saved: GeneratedReport | None = None

    async def get_active_template(self) -> ReportTemplate:
        return self.template

    async def add_generated_report(self, report: GeneratedReport) -> None:
        self.saved = report


class WeeklyDocxRendererTest(unittest.TestCase):
    def test_renders_template_content_and_fixed_table_geometry(self) -> None:
        content = WeeklyDocxRenderer().render(sample_report(), TEMPLATE_PATH)
        document = Document(BytesIO(content))
        all_text = "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
        )

        self.assertEqual(document.core_properties.title, "주간 업무보고서")
        self.assertIn("홍길동", all_text)
        self.assertIn("주간보고서 DOCX 생성 기능 개발", all_text)
        self.assertIn("실제 사용자 데이터로 주간보고서 생성 결과 확인", all_text)
        self.assertNotIn("{{", all_text)
        self.assertEqual(len(document.tables), 3)
        for table in document.tables[1:]:
            width = table._tbl.tblPr.find(qn("w:tblW"))
            layout = table._tbl.tblPr.find(qn("w:tblLayout"))
            self.assertIsNotNone(width)
            self.assertEqual(width.get(qn("w:w")), "9360")
            self.assertEqual(width.get(qn("w:type")), "dxa")
            self.assertEqual(layout.get(qn("w:type")), "fixed")


class WeeklyDocumentServiceTest(unittest.IsolatedAsyncioTestCase):
    async def test_persists_docx_and_reproducible_snapshot(self) -> None:
        repository = FakeDocumentRepository()
        with tempfile.TemporaryDirectory() as directory:
            document = await WeeklyDocumentService(repository, storage_dir=directory).generate(
                sample_report()
            )

            self.assertTrue(document.file_path.is_file())
            self.assertEqual(document.file_path.read_bytes(), document.content)
            self.assertEqual(document.filename, "weekly-report-2026-08-17-2026-08-21.docx")
            self.assertIsNotNone(repository.saved)
            assert repository.saved is not None
            self.assertEqual(repository.saved.id, document.report_id)
            self.assertEqual(repository.saved.template_id, "weekly-default-v1")
            self.assertEqual(repository.saved.content_snapshot["author"], "홍길동")
            self.assertEqual(repository.saved.status, "READY")


if __name__ == "__main__":
    unittest.main()
