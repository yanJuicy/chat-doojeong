import unittest
from io import BytesIO
from zipfile import ZipFile

from docx import Document

from app.reports.shipment import ShipmentReportRequest, ShipmentReportService
from app.reports.shipment.documents import (
    ShipmentDocumentGenerationError,
    ShipmentDocumentService,
)
from app.reports.shipment.documents.mapper import map_shipment_report
from tests.test_shipment_report_service import request_data


class ShipmentDocumentServiceTest(unittest.TestCase):
    def setUp(self) -> None:
        self.report_service = ShipmentReportService()
        self.document_service = ShipmentDocumentService(self.report_service)

    def test_mapper_formats_calculated_values_for_a_document(self) -> None:
        request = ShipmentReportRequest.model_validate(request_data())
        view = map_shipment_report(self.report_service.generate(request))

        self.assertEqual(view.metrics[0].value, "1,800 EA")
        self.assertEqual(view.metrics[1].value, "1,750 EA")
        self.assertEqual(view.metrics[-1].value, "97.2%")
        self.assertIn("출하계획 1,800 EA", view.summary)
        self.assertIn("1,750 EA를 출하", view.summary)
        self.assertEqual(view.details[1].status, "부분출하")
        self.assertEqual(view.details[1].note, "차량 지연")

    def test_generates_valid_docx_with_report_content(self) -> None:
        request = ShipmentReportRequest.model_validate(request_data())
        result = self.document_service.generate(request)

        self.assertTrue(result.content.startswith(b"PK"))
        self.assertEqual(
            result.media_type,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertEqual(
            result.filename,
            "shipment-report-2026-08-14-CUST-001.docx",
        )
        with ZipFile(BytesIO(result.content)) as archive:
            self.assertIsNone(archive.testzip())

        document = Document(BytesIO(result.content))
        paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )
        self.assertIn("2026-08-14 두정테크 출하보고서", paragraph_text)
        self.assertIn("1,800 EA", table_text)
        self.assertIn("1,750 EA", table_text)
        self.assertIn("A-1002", table_text)
        self.assertIn("차량 지연", table_text)

    def test_business_validation_failure_blocks_document_generation(self) -> None:
        data = request_data()
        data["plans"][1]["unit"] = "KG"
        data["actuals"][1]["unit"] = "KG"
        request = ShipmentReportRequest.model_validate(data)

        with self.assertRaises(ShipmentDocumentGenerationError) as context:
            self.document_service.generate(request)

        self.assertEqual(context.exception.report.status.value, "FAILED")
        self.assertTrue(
            any(item.status.value == "FAIL" for item in context.exception.report.validations)
        )


if __name__ == "__main__":
    unittest.main()
